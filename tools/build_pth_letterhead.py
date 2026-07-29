from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_LOGO = ROOT / ".docx_inspect" / "word" / "media" / "image1.png"
LOGO_CROPPED = ROOT / ".docx_inspect" / "pro_tutors_hub_logo_cropped.png"
OUT = ROOT / "Pro_Tutors_Hub_Professional_Letterhead.docx"
CONTENT_WIDTH = Inches(8.0)

PURPLE = RGBColor(48, 18, 83)
DARK_PURPLE = RGBColor(31, 12, 54)
GOLD = RGBColor(232, 166, 34)
SOFT_GOLD = RGBColor(255, 245, 219)
INK = RGBColor(34, 34, 34)
MUTED = RGBColor(102, 102, 102)
LIGHT_RULE = "D9D2E9"
GOLD_HEX = "E8A622"
PURPLE_HEX = "301253"
SOFT_GOLD_HEX = "FFF5DB"


def set_run_font(run, name="Aptos", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, attrs in kwargs.items():
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in attrs.items():
            element.set(qn("w:{}".format(key)), str(value))


def set_table_borders(table, color="FFFFFF", size="0"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def paragraph_bottom_border(paragraph, color, size=10, space=8):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)


def set_fixed_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def prepare_logo():
    if not SOURCE_LOGO.exists():
        return None
    image = Image.open(SOURCE_LOGO).convert("RGBA")
    # The source artwork is on a mostly black canvas, not true transparency.
    # Crop to the crest and key glow, then remove near-black background pixels.
    crop = (
        int(image.width * 0.055),
        int(image.height * 0.215),
        int(image.width * 0.945),
        int(image.height * 0.815),
    )
    cropped = image.crop(crop)
    pixels = cropped.load()
    for y in range(cropped.height):
        for x in range(cropped.width):
            r, g, b, a = pixels[x, y]
            if r < 16 and g < 16 and b < 16:
                pixels[x, y] = (r, g, b, 0)
    cropped.save(LOGO_CROPPED)
    return LOGO_CROPPED


def add_header(section):
    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = ""

    table = header.add_table(rows=1, cols=2, width=CONTENT_WIDTH)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=PURPLE_HEX, size="0")
    set_fixed_table_width(table, [Inches(1.55), Inches(6.45)])

    logo_cell, text_cell = table.rows[0].cells
    logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    text_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(logo_cell, PURPLE_HEX)
    set_cell_shading(text_cell, PURPLE_HEX)
    set_cell_margins(logo_cell, top=80, bottom=70, start=80, end=60)
    set_cell_margins(text_cell, top=95, bottom=75, start=120, end=100)

    p_logo = logo_cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_path = prepare_logo()
    if logo_path:
        p_logo.add_run().add_picture(str(logo_path), width=Inches(1.25))

    p = text_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    title = p.add_run("PRO TUTORS HUB")
    set_run_font(title, name="Aptos Display", size=34, color=RGBColor(255, 255, 255), bold=True)

    tag_line = text_cell.add_paragraph()
    tag_line.paragraph_format.space_before = Pt(0)
    tag_line.paragraph_format.space_after = Pt(5)
    tag = tag_line.add_run("EXCELLENCE  |  CONFIDENCE  |  SUCCESS")
    set_run_font(tag, size=12, color=GOLD, bold=True)

    contact = text_cell.add_paragraph()
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(0)
    contact.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = contact.add_run("www.protutorshub.com  |  info@protutorshub.com  |  +234 703 974 5813  |  +234 813 214 8292")
    set_run_font(run, size=10, color=RGBColor(255, 255, 255), bold=True)

    strip = header.add_table(rows=1, cols=1, width=CONTENT_WIDTH)
    strip.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(strip, color=GOLD_HEX, size="0")
    set_fixed_table_width(strip, [Inches(8.0)])
    cell = strip.rows[0].cells[0]
    set_cell_shading(cell, GOLD_HEX)
    set_cell_margins(cell, top=34, bottom=34, start=100, end=100)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    sp = cell.paragraphs[0]
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(0)
    sr = sp.add_run("Quality Tutorials  |  Exam Preparation  |  Assignment Help  |  Personalized Learning  |  Online & Physical Classes")
    set_run_font(sr, size=9.4, color=DARK_PURPLE, bold=True)


def add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    footer.paragraphs[0].text = ""
    p_rule = footer.paragraphs[0]
    p_rule.paragraph_format.space_before = Pt(0)
    p_rule.paragraph_format.space_after = Pt(3)
    paragraph_bottom_border(p_rule, PURPLE_HEX, size=14, space=1)

    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Quality Tutorials | Exam Preparation | Assignment Help | Personalized Learning | Online & Physical Classes")
    set_run_font(r, size=8.5, color=DARK_PURPLE, bold=True)

    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run("Lagos, Nigeria  |  Instagram / Facebook / Twitter: @ProTutorsHub")
    set_run_font(r2, size=8, color=MUTED)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.82)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)
    section.header_distance = Inches(0.06)
    section.footer_distance = Inches(0.10)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, PURPLE),
        ("Heading 2", 13, PURPLE),
        ("Heading 3", 12, DARK_PURPLE),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def add_meta_field(table, row_idx, label, placeholder):
    label_cell, value_cell = table.rows[row_idx].cells
    for cell in (label_cell, value_cell):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=70, bottom=55, start=80, end=80)

    label_cell.text = ""
    lp = label_cell.paragraphs[0]
    lp.paragraph_format.space_after = Pt(0)
    lr = lp.add_run(label.upper())
    set_run_font(lr, size=8.5, color=PURPLE, bold=True)
    set_cell_shading(label_cell, SOFT_GOLD_HEX)
    set_cell_borders(label_cell, bottom={"val": "single", "sz": "4", "color": LIGHT_RULE})

    value_cell.text = ""
    vp = value_cell.paragraphs[0]
    vp.paragraph_format.space_after = Pt(0)
    vr = vp.add_run(placeholder)
    set_run_font(vr, size=10.5, color=MUTED)
    set_cell_borders(value_cell, bottom={"val": "single", "sz": "4", "color": LIGHT_RULE})


def add_writing_line(doc, after=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    paragraph_bottom_border(p, "C9C9C9", size=4, space=1)


def build():
    doc = Document()
    style_document(doc)
    section = doc.sections[0]
    add_header(section)
    add_footer(section)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(10)
    intro.paragraph_format.space_after = Pt(8)
    intro.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = intro.add_run("OFFICIAL CORRESPONDENCE")
    set_run_font(r, name="Aptos Display", size=13.5, color=PURPLE, bold=True)

    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(meta, color="FFFFFF", size="0")
    set_fixed_table_width(meta, [Inches(1.35), Inches(6.65)])
    add_meta_field(meta, 0, "Date", "June 3, 2026")
    add_meta_field(meta, 1, "To", "Recipient name / organization")
    add_meta_field(meta, 2, "Subject", "Subject of correspondence")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)

    salutation = doc.add_paragraph()
    salutation.paragraph_format.space_after = Pt(8)
    sr = salutation.add_run("Dear ______________________________,")
    set_run_font(sr, size=11, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(
        "Use this space for your official message. The new layout keeps the Pro Tutors Hub name prominent while leaving a clear writing area for formal correspondence."
    )
    set_run_font(r, size=10.8, color=INK)

    for _ in range(8):
        add_writing_line(doc)

    signoff = doc.add_paragraph()
    signoff.paragraph_format.space_before = Pt(14)
    signoff.paragraph_format.space_after = Pt(20)
    r = signoff.add_run("Yours faithfully,")
    set_run_font(r, size=11, color=INK)

    sig_table = doc.add_table(rows=1, cols=2)
    set_table_borders(sig_table, color="FFFFFF", size="0")
    set_fixed_table_width(sig_table, [Inches(3.8), Inches(4.2)])
    left, right = sig_table.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, top=40, bottom=40, start=0, end=80)
    left.text = ""
    p = left.paragraphs[0]
    paragraph_bottom_border(p, LIGHT_RULE, size=6, space=2)
    p.paragraph_format.space_after = Pt(3)
    label = left.add_paragraph()
    label.paragraph_format.space_after = Pt(0)
    lr = label.add_run("Authorized Signature")
    set_run_font(lr, size=8.5, color=MUTED)

    right.text = ""
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run("PRO TUTORS HUB")
    set_run_font(rr, name="Aptos Display", size=11, color=PURPLE, bold=True)
    rt = right.add_paragraph()
    rt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tr = rt.add_run("Excellence | Confidence | Success")
    set_run_font(tr, size=8.5, color=GOLD, bold=True)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
